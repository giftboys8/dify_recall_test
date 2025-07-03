"""PDF解析器模块

基于pdf2docx和docx库实现PDF文档的解析和文本提取。
"""

import os
import re
import tempfile
from typing import List, Dict, Any, Optional
from pathlib import Path

try:
    from pdf2docx import Converter
    import docx
except ImportError as e:
    raise ImportError(f"缺少必要的依赖包: {e}. 请安装: pip install pdf2docx python-docx")

from ..utils.logger import get_logger

logger = get_logger(__name__)


class PDFParser:
    """PDF文档解析器
    
    支持PDF转换为结构化文档，提取文本内容并保持格式信息。
    """
    
    def __init__(self, temp_dir: Optional[str] = None):
        """初始化PDF解析器
        
        Args:
            temp_dir: 临时文件目录，默认使用系统临时目录
        """
        self.temp_dir = temp_dir or tempfile.gettempdir()
        self.temp_files = []  # 跟踪临时文件用于清理
        
    def parse_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """解析PDF文档
        
        Args:
            pdf_path: PDF文件路径
            
        Returns:
            包含文档结构和文本内容的字典
        """
        try:
            logger.info(f"开始解析PDF文档: {pdf_path}")
            
            # 验证文件存在
            if not os.path.exists(pdf_path):
                raise FileNotFoundError(f"PDF文件不存在: {pdf_path}")
                
            # 生成临时DOCX文件路径
            temp_docx = self._get_temp_file_path("temp.docx")
            
            # Step1: PDF转结构化DOCX
            cv = Converter(pdf_path)
            cv.convert(temp_docx, keep_layout=True)
            cv.close()
            
            logger.info(f"PDF转换完成: {temp_docx}")
            
            # Step2: 提取文档结构和文本
            doc_data = self._extract_document_data(temp_docx)
            
            return {
                'success': True,
                'temp_docx_path': temp_docx,
                'document_data': doc_data,
                'source_pdf': pdf_path
            }
            
        except Exception as e:
            logger.error(f"PDF解析失败: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'source_pdf': pdf_path
            }
    
    def _extract_document_data(self, docx_path: str) -> Dict[str, Any]:
        """提取DOCX文档的结构和文本数据
        
        Args:
            docx_path: DOCX文件路径
            
        Returns:
            文档数据字典
        """
        doc = docx.Document(docx_path)
        
        # 提取段落文本
        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():  # 跳过空段落
                paragraphs.append({
                    'index': i,
                    'text': para.text,
                    'style': para.style.name if para.style else None
                })
        
        # 提取表格数据
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append({
                'index': i,
                'data': table_data
            })
        
        return {
            'paragraphs': paragraphs,
            'tables': tables,
            'total_paragraphs': len(paragraphs),
            'total_tables': len(tables)
        }
    
    def get_text_for_translation(self, doc_data: Dict[str, Any], use_smart_chunking: bool = True, max_chars: int = 1500, min_chars: int = 50) -> List[str]:
        """获取需要翻译的文本列表
        
        Args:
            doc_data: 文档数据
            use_smart_chunking: 是否使用智能分块
            
        Returns:
            文本列表
        """
        texts = []
        
        # 添加段落文本
        for para in doc_data['paragraphs']:
            texts.append(para['text'])
        
        # 添加表格文本
        for table in doc_data['tables']:
            for row in table['data']:
                for cell in row:
                    if cell.strip():
                        texts.append(cell)
        
        # 应用智能分块
        if use_smart_chunking:
            texts = self._apply_smart_chunking(texts, max_chars, min_chars)
        
        return texts
    
    def _apply_smart_chunking(self, texts: List[str], max_chars: int = 1500, min_chars: int = 50) -> List[str]:
        """应用智能文本分块
        
        Args:
            texts: 原始文本列表
            max_chars: 最大字符数
            min_chars: 最小字符数
            
        Returns:
            优化后的文本列表
        """
        chunked_texts = []
        current_chunk = ""
        
        for text in texts:
            text = text.strip()
            if not text:
                continue
            
            # 如果文本过长，需要分割
            if len(text) > max_chars:
                # 先保存当前累积的块
                if current_chunk:
                    chunked_texts.append(current_chunk.strip())
                    current_chunk = ""
                
                # 分割长文本
                split_texts = self._split_long_text(text, max_chars)
                chunked_texts.extend(split_texts)
            
            # 如果文本较短，尝试与其他短文本合并
            elif len(text) < min_chars:
                if len(current_chunk + " " + text) <= max_chars:
                    current_chunk += (" " if current_chunk else "") + text
                else:
                    if current_chunk:
                        chunked_texts.append(current_chunk.strip())
                    current_chunk = text
            
            # 中等长度文本
            else:
                # 先保存当前累积的块
                if current_chunk:
                    chunked_texts.append(current_chunk.strip())
                    current_chunk = ""
                
                chunked_texts.append(text)
        
        # 保存最后的累积块
        if current_chunk:
            chunked_texts.append(current_chunk.strip())
        
        logger.info(f"智能分块: {len(texts)} -> {len(chunked_texts)} 个文本块")
        return chunked_texts
    
    def _split_long_text(self, text: str, max_chars: int) -> List[str]:
        """分割长文本，保持语义完整性
        
        Args:
            text: 要分割的文本
            max_chars: 最大字符数
            
        Returns:
            分割后的文本列表
        """
        if len(text) <= max_chars:
            return [text]
        
        chunks = []
        
        # 首先尝试按句子分割（中文和英文句号）
        sentences = re.split(r'[。！？.!?]\s*', text)
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # 恢复句号
            if not sentence.endswith(('。', '！', '？', '.', '!', '?')):
                sentence += '。'
            
            # 检查是否可以添加到当前块
            test_chunk = current_chunk + (" " if current_chunk else "") + sentence
            
            if len(test_chunk) <= max_chars:
                current_chunk = test_chunk
            else:
                # 保存当前块
                if current_chunk:
                    chunks.append(current_chunk)
                
                # 如果单个句子仍然太长，按字符强制分割
                if len(sentence) > max_chars:
                    chunks.extend(self._force_split_text(sentence, max_chars))
                    current_chunk = ""
                else:
                    current_chunk = sentence
        
        # 保存最后的块
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def _force_split_text(self, text: str, max_chars: int) -> List[str]:
        """强制按字符数分割文本
        
        Args:
            text: 要分割的文本
            max_chars: 最大字符数
            
        Returns:
            分割后的文本列表
        """
        chunks = []
        for i in range(0, len(text), max_chars):
            chunk = text[i:i + max_chars]
            if chunk.strip():
                chunks.append(chunk.strip())
        return chunks
    
    def _get_temp_file_path(self, filename: str) -> str:
        """生成临时文件路径
        
        Args:
            filename: 文件名
            
        Returns:
            临时文件完整路径
        """
        temp_path = os.path.join(self.temp_dir, filename)
        self.temp_files.append(temp_path)
        return temp_path
    
    def cleanup(self):
        """清理临时文件"""
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
                    logger.debug(f"已删除临时文件: {temp_file}")
            except Exception as e:
                logger.warning(f"删除临时文件失败 {temp_file}: {e}")
        
        self.temp_files.clear()
    
    def __del__(self):
        """析构函数，自动清理临时文件"""
        self.cleanup()