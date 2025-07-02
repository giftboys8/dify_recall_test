"""文档格式化模块

处理DOCX文档的结构化操作和PDF转换功能。
"""

import os
import tempfile
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

try:
    from docx2pdf import convert as docx2pdf_convert
except ImportError:
    docx2pdf_convert = None

from ..utils.logger import get_logger

logger = get_logger(__name__)


class DocumentFormatter:
    """文档格式化器"""
    
    def __init__(self):
        self.logger = get_logger(__name__)
        self._check_dependencies()
    
    def _check_dependencies(self):
        """检查依赖库"""
        if Document is None:
            raise ImportError("python-docx库未安装，请运行: pip install python-docx")
    
    def create_bilingual_document(self, 
                                original_texts: List[str], 
                                translated_texts: List[str],
                                output_path: str,
                                title: str = "双语对照文档",
                                layout: str = "side_by_side") -> bool:
        """创建双语对照文档
        
        Args:
            original_texts: 原文文本列表
            translated_texts: 译文文本列表
            output_path: 输出文件路径
            title: 文档标题
            layout: 布局方式 ('side_by_side', 'paragraph_by_paragraph')
            
        Returns:
            是否创建成功
        """
        try:
            if len(original_texts) != len(translated_texts):
                raise ValueError("原文和译文数量不匹配")
            
            doc = Document()
            
            # 添加标题
            title_paragraph = doc.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if layout == "side_by_side":
                self._create_side_by_side_layout(doc, original_texts, translated_texts)
            else:
                self._create_paragraph_layout(doc, original_texts, translated_texts)
            
            # 保存文档
            doc.save(output_path)
            self.logger.info(f"双语文档已保存: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"创建双语文档失败: {e}")
            return False
    
    def _create_side_by_side_layout(self, doc: Document, 
                                   original_texts: List[str], 
                                   translated_texts: List[str]):
        """创建左右对照布局"""
        # 创建表格
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # 设置表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '原文'
        header_cells[1].text = '译文'
        
        # 添加内容行
        for original, translated in zip(original_texts, translated_texts):
            if original.strip() or translated.strip():  # 跳过空行
                row_cells = table.add_row().cells
                row_cells[0].text = original
                row_cells[1].text = translated
    
    def _create_paragraph_layout(self, doc: Document, 
                               original_texts: List[str], 
                               translated_texts: List[str]):
        """创建段落对照布局"""
        for i, (original, translated) in enumerate(zip(original_texts, translated_texts)):
            if original.strip() or translated.strip():  # 跳过空行
                # 添加段落编号
                doc.add_heading(f'段落 {i+1}', level=2)
                
                # 原文
                original_para = doc.add_paragraph()
                original_para.add_run('原文: ').bold = True
                original_para.add_run(original)
                
                # 译文
                translated_para = doc.add_paragraph()
                translated_para.add_run('译文: ').bold = True
                translated_para.add_run(translated)
                
                # 添加分隔线
                doc.add_paragraph('─' * 50)
    
    def merge_translations_to_docx(self, 
                                 docx_path: str, 
                                 translations: List[str],
                                 output_path: str,
                                 replace_original: bool = False) -> bool:
        """将翻译结果合并到DOCX文档
        
        Args:
            docx_path: 原始DOCX文件路径
            translations: 翻译文本列表
            output_path: 输出文件路径
            replace_original: 是否替换原文（True）或创建双语文档（False）
            
        Returns:
            是否合并成功
        """
        try:
            doc = Document(docx_path)
            
            # 提取所有段落
            paragraphs = [p for p in doc.paragraphs if p.text.strip()]
            
            if len(paragraphs) != len(translations):
                self.logger.warning(f"段落数量({len(paragraphs)})与翻译数量({len(translations)})不匹配")
            
            if replace_original:
                # 替换原文
                for i, translation in enumerate(translations):
                    if i < len(paragraphs):
                        paragraphs[i].text = translation
            else:
                # 创建双语文档
                for i, translation in enumerate(translations):
                    if i < len(paragraphs):
                        original_text = paragraphs[i].text
                        paragraphs[i].text = f"{original_text}\n\n[译文] {translation}"
            
            # 保存文档
            doc.save(output_path)
            self.logger.info(f"翻译已合并到文档: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"合并翻译失败: {e}")
            return False
    
    def convert_docx_to_pdf(self, docx_path: str, pdf_path: str) -> bool:
        """将DOCX转换为PDF
        
        Args:
            docx_path: DOCX文件路径
            pdf_path: PDF输出路径
            
        Returns:
            是否转换成功
        """
        try:
            if docx2pdf_convert is None:
                self.logger.error("docx2pdf库未安装，无法转换为PDF")
                return False
            
            # 确保输出目录存在
            os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
            
            # 执行转换
            docx2pdf_convert(docx_path, pdf_path)
            
            self.logger.info(f"PDF转换完成: {pdf_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"DOCX转PDF失败: {e}")
            return False
    
    def extract_text_from_docx(self, docx_path: str) -> List[str]:
        """从DOCX文档提取文本
        
        Args:
            docx_path: DOCX文件路径
            
        Returns:
            文本段落列表
        """
        try:
            doc = Document(docx_path)
            texts = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # 跳过空段落
                    texts.append(text)
            
            # 处理表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            texts.append(text)
            
            self.logger.info(f"从DOCX提取了 {len(texts)} 个文本段落")
            return texts
            
        except Exception as e:
            self.logger.error(f"从DOCX提取文本失败: {e}")
            return []
    
    def create_translation_report(self, 
                                translation_result: Dict[str, Any],
                                output_path: str) -> bool:
        """创建翻译报告
        
        Args:
            translation_result: 翻译结果字典
            output_path: 报告输出路径
            
        Returns:
            是否创建成功
        """
        try:
            doc = Document()
            
            # 标题
            title = doc.add_heading('PDF翻译报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 基本信息
            doc.add_heading('翻译信息', level=1)
            info_table = doc.add_table(rows=6, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ('翻译状态', '成功' if translation_result.get('success') else '失败'),
                ('翻译提供商', translation_result.get('provider', 'N/A')),
                ('原文段落数', str(translation_result.get('original_count', 0))),
                ('译文段落数', str(translation_result.get('translated_count', 0))),
                ('翻译耗时', f"{translation_result.get('duration', 0):.2f}秒"),
                ('翻译时间', translation_result.get('timestamp', 'N/A'))
            ]
            
            for i, (key, value) in enumerate(info_data):
                info_table.cell(i, 0).text = key
                info_table.cell(i, 1).text = value
            
            # 错误信息（如果有）
            if not translation_result.get('success') and translation_result.get('error'):
                doc.add_heading('错误信息', level=1)
                doc.add_paragraph(translation_result['error'])
            
            # 翻译统计
            if translation_result.get('success'):
                doc.add_heading('翻译统计', level=1)
                
                original_texts = translation_result.get('original_texts', [])
                translated_texts = translation_result.get('translated_texts', [])
                
                if original_texts and translated_texts:
                    # 字符统计
                    original_chars = sum(len(text) for text in original_texts)
                    translated_chars = sum(len(text) for text in translated_texts)
                    
                    stats_para = doc.add_paragraph()
                    stats_para.add_run(f"原文总字符数: {original_chars}\n")
                    stats_para.add_run(f"译文总字符数: {translated_chars}\n")
                    stats_para.add_run(f"平均翻译比率: {translated_chars/original_chars:.2f}" if original_chars > 0 else "平均翻译比率: N/A")
            
            # 保存报告
            doc.save(output_path)
            self.logger.info(f"翻译报告已保存: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"创建翻译报告失败: {e}")
            return False
    
    @staticmethod
    def is_docx2pdf_available() -> bool:
        """检查docx2pdf是否可用"""
        return docx2pdf_convert is not None
    
    @staticmethod
    def get_supported_formats() -> List[str]:
        """获取支持的输出格式"""
        formats = ['docx']
        if DocumentFormatter.is_docx2pdf_available():
            formats.append('pdf')
        return formats